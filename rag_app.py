"""
RAG (Retrieval-Augmented Generation) Application
==================================================
A full-featured RAG system with:
- Document processing (PDF, HTML, DOCX, TXT, MD)
- Vector database (ChromaDB with persistent storage)
- Hybrid search (semantic + BM25 keyword search)
- Conversation memory (last 10 exchanges)
- Streaming LLM responses with source citations
- Gradio-based conversational UI

Requirements (install via pip):
    pip install chromadb sentence-transformers gradio openai pymupdf python-docx \
                beautifulsoup4 rank_bm25 nltk tiktoken numpy

Usage:
    1. Place 50+ documents in a ./documents/ folder (PDF, HTML, DOCX, TXT, MD)
    2. Set your OpenAI API key:  export OPENAI_API_KEY="sk-..."
    3. Run:  python rag_app.py
    4. Open the Gradio URL in your browser
"""

import os
import re
import json
import hashlib
import logging
import textwrap
from pathlib import Path
from typing import Optional
from dataclasses import dataclass, field
from collections import defaultdict

import numpy as np

# -- Document parsing --
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from bs4 import BeautifulSoup

# -- NLP / chunking --
import nltk
from nltk.tokenize import sent_tokenize

# -- Embeddings & vector DB --
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.config import Settings

# -- BM25 keyword search --
from rank_bm25 import BM25Okapi

# -- LLM --
import openai

# -- UI --
import gradio as gr

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

@dataclass
class Config:
    """Central configuration for the RAG pipeline."""

    # Paths
    documents_dir: str = "./documents"
    chroma_persist_dir: str = "./chroma_db"

    # Chunking
    chunk_size: int = 512          # target tokens per chunk (sentence-based)
    chunk_overlap: int = 64        # overlap tokens between consecutive chunks
    min_chunk_length: int = 40     # discard chunks shorter than this (chars)

    # Embedding model (runs locally via sentence-transformers)
    embedding_model: str = "all-MiniLM-L6-v2"
    chroma_collection: str = "rag_docs"

    # Retrieval
    top_k_semantic: int = 20       # initial semantic retrieval
    top_k_bm25: int = 20           # initial BM25 retrieval
    top_k_final: int = 5           # after hybrid merge / re-rank

    # Hybrid search weight  (0 = pure BM25, 1 = pure semantic)
    semantic_weight: float = 0.6

    # LLM
    openai_model: str = "gpt-4o-mini"
    temperature: float = 0.2
    max_context_tokens: int = 6000
    system_prompt: str = textwrap.dedent("""\
        You are a knowledgeable assistant. Answer the user's question using ONLY
        the provided context passages. If the context does not contain enough
        information, say so honestly.

        Rules:
        - Cite sources using [Source N] notation after each claim.
        - Be concise but thorough.
        - If multiple sources agree, prefer the most specific one.
        - For follow-up questions, use conversation history for context.
    """)

    # Conversation memory
    memory_length: int = 10  # number of past exchanges to keep

    # Server
    server_port: int = 7860
    share: bool = True  # set True for public URL via Gradio


CFG = Config()

# ---------------------------------------------------------------------------
# Logging
# ---------------------------------------------------------------------------

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-7s | %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger("rag")

# ---------------------------------------------------------------------------
# 1. Document Processing
# ---------------------------------------------------------------------------

@dataclass
class RawDocument:
    """A single extracted document before chunking."""
    text: str
    metadata: dict = field(default_factory=dict)


def extract_pdf(path: str) -> RawDocument:
    """Extract text and metadata from a PDF using PyMuPDF."""
    doc = fitz.open(path)
    pages = []
    for page in doc:
        pages.append(page.get_text("text"))
    meta = doc.metadata or {}
    return RawDocument(
        text="\n\n".join(pages),
        metadata={
            "source": os.path.basename(path),
            "path": path,
            "type": "pdf",
            "title": meta.get("title", ""),
            "author": meta.get("author", ""),
            "pages": len(doc),
        },
    )


def extract_docx(path: str) -> RawDocument:
    """Extract text from a DOCX file."""
    doc = DocxDocument(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    core = doc.core_properties
    return RawDocument(
        text="\n\n".join(paragraphs),
        metadata={
            "source": os.path.basename(path),
            "path": path,
            "type": "docx",
            "title": core.title or "",
            "author": core.author or "",
        },
    )


def extract_html(path: str) -> RawDocument:
    """Extract text from an HTML file."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")
    # Remove script and style elements
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    title = soup.title.string if soup.title else ""
    text = soup.get_text(separator="\n", strip=True)
    return RawDocument(
        text=text,
        metadata={
            "source": os.path.basename(path),
            "path": path,
            "type": "html",
            "title": title,
        },
    )


def extract_text(path: str) -> RawDocument:
    """Extract text from a plain text or markdown file."""
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        text = f.read()
    return RawDocument(
        text=text,
        metadata={
            "source": os.path.basename(path),
            "path": path,
            "type": "text",
        },
    )


EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".html": extract_html,
    ".htm": extract_html,
    ".txt": extract_text,
    ".md": extract_text,
}


def load_documents(directory: str) -> list[RawDocument]:
    """Recursively load all supported documents from a directory."""
    docs = []
    directory = Path(directory)
    if not directory.exists():
        log.warning(f"Documents directory not found: {directory}")
        return docs

    for fpath in sorted(directory.rglob("*")):
        ext = fpath.suffix.lower()
        if ext in EXTRACTORS:
            try:
                doc = EXTRACTORS[ext](str(fpath))
                if len(doc.text.strip()) > 50:
                    docs.append(doc)
                    log.info(f"  Loaded: {fpath.name} ({len(doc.text):,} chars)")
            except Exception as e:
                log.error(f"  Failed: {fpath.name} -> {e}")
    log.info(f"Total documents loaded: {len(docs)}")
    return docs

# ---------------------------------------------------------------------------
# 2. Smart Chunking (Sentence-Based with Overlap)
# ---------------------------------------------------------------------------

@dataclass
class Chunk:
    """A text chunk ready for embedding."""
    text: str
    metadata: dict
    chunk_id: str


def _approx_token_count(text: str) -> int:
    """Rough token count (≈ 4 chars per token for English)."""
    return len(text) // 4


def sentence_chunk(doc: RawDocument, chunk_size: int = 512, overlap: int = 64) -> list[Chunk]:
    """
    Sentence-based chunking strategy:
    - Split text into sentences.
    - Accumulate sentences until chunk_size tokens is reached.
    - Overlap by re-including trailing sentences from previous chunk.
    """
    try:
        sentences = sent_tokenize(doc.text)
    except Exception:
        nltk.download("punkt_tab", quiet=True)
        sentences = sent_tokenize(doc.text)

    if not sentences:
        return []

    chunks: list[Chunk] = []
    current_sentences: list[str] = []
    current_tokens = 0

    def _flush(sents: list[str], idx: int):
        text = " ".join(sents).strip()
        if len(text) < CFG.min_chunk_length:
            return
        chunk_id = hashlib.md5(
            f"{doc.metadata.get('source', '')}:{idx}:{text[:80]}".encode()
        ).hexdigest()[:12]
        chunks.append(Chunk(
            text=text,
            metadata={**doc.metadata, "chunk_index": idx},
            chunk_id=chunk_id,
        ))

    chunk_idx = 0
    for sent in sentences:
        sent_tokens = _approx_token_count(sent)
        if current_tokens + sent_tokens > chunk_size and current_sentences:
            _flush(current_sentences, chunk_idx)
            chunk_idx += 1
            # Keep overlap sentences from the tail
            overlap_sents: list[str] = []
            overlap_tok = 0
            for s in reversed(current_sentences):
                t = _approx_token_count(s)
                if overlap_tok + t > overlap:
                    break
                overlap_sents.insert(0, s)
                overlap_tok += t
            current_sentences = overlap_sents
            current_tokens = overlap_tok
        current_sentences.append(sent)
        current_tokens += sent_tokens

    if current_sentences:
        _flush(current_sentences, chunk_idx)

    return chunks


def chunk_all_documents(docs: list[RawDocument]) -> list[Chunk]:
    """Chunk every loaded document."""
    all_chunks = []
    for doc in docs:
        doc_chunks = sentence_chunk(doc, CFG.chunk_size, CFG.chunk_overlap)
        all_chunks.extend(doc_chunks)
    log.info(f"Total chunks created: {len(all_chunks)}")
    return all_chunks

# ---------------------------------------------------------------------------
# 3. Vector Database (ChromaDB with Persistent Storage)
# ---------------------------------------------------------------------------

class VectorStore:
    """Manages ChromaDB collection and embedding model."""

    def __init__(self, config: Config):
        self.config = config
        log.info(f"Loading embedding model: {config.embedding_model}")
        self.embedder = SentenceTransformer(config.embedding_model)

        self.client = chromadb.Client(Settings(
            persist_directory=config.chroma_persist_dir,
            anonymized_telemetry=False,
            is_persistent=True,
        ))
        self.collection = self.client.get_or_create_collection(
            name=config.chroma_collection,
            metadata={"hnsw:space": "cosine"},
        )
        log.info(
            f"ChromaDB collection '{config.chroma_collection}' "
            f"has {self.collection.count()} vectors"
        )

    def embed_text(self, texts: list[str]) -> list[list[float]]:
        """Generate embeddings for a list of texts."""
        return self.embedder.encode(texts, show_progress_bar=False).tolist()

    def embed_single(self, text: str) -> list[float]:
        """Embed a single query string."""
        return self.embedder.encode(text).tolist()

    def add_chunks(self, chunks: list[Chunk], batch_size: int = 256):
        """Insert chunks into ChromaDB (skip duplicates by ID)."""
        existing = set(self.collection.get()["ids"]) if self.collection.count() > 0 else set()
        new_chunks = [c for c in chunks if c.chunk_id not in existing]
        if not new_chunks:
            log.info("No new chunks to add (all already indexed).")
            return

        for i in range(0, len(new_chunks), batch_size):
            batch = new_chunks[i : i + batch_size]
            ids = [c.chunk_id for c in batch]
            texts = [c.text for c in batch]
            metas = [c.metadata for c in batch]
            embeddings = self.embed_text(texts)
            self.collection.add(
                ids=ids,
                documents=texts,
                metadatas=metas,
                embeddings=embeddings,
            )
            log.info(f"  Indexed batch {i // batch_size + 1} ({len(batch)} chunks)")

        log.info(f"Total vectors in DB: {self.collection.count()}")

    def semantic_search(self, query: str, k: int = 20) -> list[dict]:
        """Return top-k results by cosine similarity."""
        count = self.collection.count()
        if count == 0:
            return []
        embedding = self.embed_single(query)
        results = self.collection.query(
            query_embeddings=[embedding],
            n_results=min(k, count),
            include=["documents", "metadatas", "distances"],
        )
        hits = []
        if results["documents"] and results["documents"][0]:
            for doc, meta, dist in zip(
                results["documents"][0],
                results["metadatas"][0],
                results["distances"][0],
            ):
                hits.append({
                    "text": doc,
                    "metadata": meta,
                    "score": 1 - dist,  # cosine distance -> similarity
                })
        return hits

# ---------------------------------------------------------------------------
# 4. BM25 Keyword Search (for Hybrid Retrieval)
# ---------------------------------------------------------------------------

class BM25Index:
    """Maintains a BM25 index over all chunk texts."""

    def __init__(self):
        self.corpus: list[str] = []
        self.metadata: list[dict] = []
        self.bm25: Optional[BM25Okapi] = None

    def build(self, chunks: list[Chunk]):
        """Build BM25 index from chunks."""
        self.corpus = [c.text for c in chunks]
        self.metadata = [c.metadata for c in chunks]
        tokenized = [self._tokenize(t) for t in self.corpus]
        self.bm25 = BM25Okapi(tokenized)
        log.info(f"BM25 index built over {len(self.corpus)} chunks")

    @staticmethod
    def _tokenize(text: str) -> list[str]:
        return re.findall(r"\w+", text.lower())

    def search(self, query: str, k: int = 20) -> list[dict]:
        """Return top-k BM25 results."""
        if self.bm25 is None:
            return []
        tokens = self._tokenize(query)
        scores = self.bm25.get_scores(tokens)
        top_idx = np.argsort(scores)[::-1][:k]
        results = []
        for idx in top_idx:
            if scores[idx] > 0:
                results.append({
                    "text": self.corpus[idx],
                    "metadata": self.metadata[idx],
                    "score": float(scores[idx]),
                })
        return results

# ---------------------------------------------------------------------------
# 5. Hybrid Search: Merge Semantic + BM25 with RRF
# ---------------------------------------------------------------------------

def reciprocal_rank_fusion(
    semantic_hits: list[dict],
    bm25_hits: list[dict],
    semantic_weight: float = 0.6,
    k_constant: int = 60,
    top_k: int = 5,
) -> list[dict]:
    """
    Reciprocal Rank Fusion (RRF) to merge two ranked lists.
    score(doc) = w_s / (k + rank_semantic) + w_b / (k + rank_bm25)
    """
    scores: dict[str, float] = defaultdict(float)
    doc_map: dict[str, dict] = {}

    bm25_weight = 1.0 - semantic_weight

    for rank, hit in enumerate(semantic_hits, start=1):
        key = hit["text"][:200]  # use text prefix as dedup key
        scores[key] += semantic_weight / (k_constant + rank)
        doc_map[key] = hit

    for rank, hit in enumerate(bm25_hits, start=1):
        key = hit["text"][:200]
        scores[key] += bm25_weight / (k_constant + rank)
        if key not in doc_map:
            doc_map[key] = hit

    ranked = sorted(scores.items(), key=lambda x: x[1], reverse=True)[:top_k]
    results = []
    for key, score in ranked:
        entry = doc_map[key].copy()
        entry["hybrid_score"] = score
        results.append(entry)
    return results

# ---------------------------------------------------------------------------
# 6. Conversation Memory
# ---------------------------------------------------------------------------

class ConversationMemory:
    """Tracks the last N exchanges for multi-turn support."""

    def __init__(self, max_turns: int = 10):
        self.max_turns = max_turns
        self.history: list[dict] = []  # [{"role": "user"/"assistant", "content": ...}]

    def add_user(self, message: str):
        self.history.append({"role": "user", "content": message})
        self._trim()

    def add_assistant(self, message: str):
        self.history.append({"role": "assistant", "content": message})
        self._trim()

    def _trim(self):
        # Keep last N *exchanges* (each exchange = 2 messages)
        max_messages = self.max_turns * 2
        if len(self.history) > max_messages:
            self.history = self.history[-max_messages:]

    def get_messages(self) -> list[dict]:
        return list(self.history)

    def get_context_summary(self) -> str:
        """Produce a short summary for query rewriting."""
        if not self.history:
            return ""
        recent = self.history[-6:]  # last 3 exchanges
        lines = []
        for msg in recent:
            role = "User" if msg["role"] == "user" else "Assistant"
            # Truncate long assistant replies
            content = msg["content"][:300]
            lines.append(f"{role}: {content}")
        return "\n".join(lines)

    def clear(self):
        self.history.clear()

# ---------------------------------------------------------------------------
# 7. RAG Pipeline (Query → Retrieve → Generate)
# ---------------------------------------------------------------------------

class RAGPipeline:
    """Orchestrates the full RAG pipeline."""

    def __init__(self, config: Config):
        self.config = config
        self.vector_store = VectorStore(config)
        self.bm25_index = BM25Index()
        self.memory = ConversationMemory(max_turns=config.memory_length)
        self.openai_client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY", ""))

    # -- Indexing ----------------------------------------------------------

    def index_documents(self, docs_dir: Optional[str] = None):
        """Load, chunk, and index all documents."""
        directory = docs_dir or self.config.documents_dir
        raw_docs = load_documents(directory)
        if not raw_docs:
            log.warning("No documents found. Please add files to the documents folder.")
            return

        chunks = chunk_all_documents(raw_docs)
        self.vector_store.add_chunks(chunks)
        self.bm25_index.build(chunks)
        log.info("Indexing complete.")

    # -- Query rewriting for follow-ups ------------------------------------

    def _rewrite_query(self, user_query: str) -> str:
        """Use conversation context to make follow-up queries self-contained."""
        context = self.memory.get_context_summary()
        if not context:
            return user_query

        try:
            response = self.openai_client.chat.completions.create(
                model=self.config.openai_model,
                temperature=0,
                max_tokens=200,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "Rewrite the user's latest question so it is self-contained, "
                            "incorporating any necessary context from the conversation. "
                            "Output ONLY the rewritten question, nothing else."
                        ),
                    },
                    {
                        "role": "user",
                        "content": f"Conversation:\n{context}\n\nLatest question: {user_query}",
                    },
                ],
            )
            rewritten = response.choices[0].message.content.strip()
            if rewritten:
                log.info(f"Rewritten query: {rewritten}")
                return rewritten
        except Exception as e:
            log.warning(f"Query rewriting failed: {e}")
        return user_query

    # -- Retrieval ---------------------------------------------------------

    def retrieve(self, query: str) -> list[dict]:
        """Hybrid retrieval: semantic + BM25 merged via RRF."""
        semantic_hits = self.vector_store.semantic_search(
            query, k=self.config.top_k_semantic
        )
        bm25_hits = self.bm25_index.search(query, k=self.config.top_k_bm25)

        merged = reciprocal_rank_fusion(
            semantic_hits,
            bm25_hits,
            semantic_weight=self.config.semantic_weight,
            top_k=self.config.top_k_final,
        )
        return merged

    # -- Generation --------------------------------------------------------

    def generate(self, user_query: str, retrieved_chunks: list[dict]) -> str:
        """Call the LLM with retrieved context and conversation history."""

        # Build context block with source labels
        context_parts = []
        for i, chunk in enumerate(retrieved_chunks, 1):
            source = chunk["metadata"].get("source", "unknown")
            title = chunk["metadata"].get("title", "")
            label = f"[Source {i}: {source}"
            if title:
                label += f" — {title}"
            label += "]"
            context_parts.append(f"{label}\n{chunk['text']}")

        context_block = "\n\n---\n\n".join(context_parts)

        # Assemble messages
        messages = [{"role": "system", "content": self.config.system_prompt}]

        # Add conversation history
        messages.extend(self.memory.get_messages())

        # Add current turn with context
        user_content = (
            f"Context passages:\n\n{context_block}\n\n"
            f"---\n\nQuestion: {user_query}"
        )
        messages.append({"role": "user", "content": user_content})

        try:
            response = self.openai_client.chat.completions.create(
                model=self.config.openai_model,
                temperature=self.config.temperature,
                max_tokens=1500,
                messages=messages,
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"LLM generation error: {e}"

    # -- Full pipeline -----------------------------------------------------

    def query(self, user_input: str) -> tuple[str, list[dict]]:
        """
        Full RAG pipeline:
        1. Rewrite query using conversation context
        2. Hybrid retrieve top-K chunks
        3. Generate answer with citations
        4. Update memory
        Returns (answer_text, retrieved_sources)
        """
        # Step 1: Rewrite for follow-ups
        search_query = self._rewrite_query(user_input)

        # Step 2: Retrieve
        chunks = self.retrieve(search_query)
        if not chunks:
            answer = (
                "I couldn't find any relevant information in the document collection "
                "to answer your question. Could you rephrase or ask about a different topic?"
            )
            self.memory.add_user(user_input)
            self.memory.add_assistant(answer)
            return answer, []

        # Step 3: Generate
        answer = self.generate(user_input, chunks)

        # Step 4: Update memory
        self.memory.add_user(user_input)
        self.memory.add_assistant(answer)

        return answer, chunks

    def reset_conversation(self):
        """Clear conversation history."""
        self.memory.clear()
        return "Conversation history cleared."

# ---------------------------------------------------------------------------
# 8. Gradio Conversational UI
# ---------------------------------------------------------------------------

def build_ui(pipeline: RAGPipeline) -> gr.Blocks:
    """Create the Gradio chat interface with source citations."""

    CUSTOM_CSS = """
    .gradio-container {
        max-width: 960px !important;
        margin: auto !important;
        font-family: 'Segoe UI', system-ui, sans-serif !important;
    }
    .source-card {
        background: #f8f9fa;
        border-left: 3px solid #4a90d9;
        padding: 10px 14px;
        margin: 6px 0;
        border-radius: 4px;
        font-size: 0.88em;
        line-height: 1.5;
    }
    .source-card strong { color: #2c5282; }
    .status-bar {
        text-align: center;
        padding: 6px;
        font-size: 0.85em;
        color: #718096;
    }
    """

    with gr.Blocks(css=CUSTOM_CSS, title="RAG Assistant", theme=gr.themes.Soft()) as demo:
        gr.Markdown(
            "# 📚 RAG Document Assistant\n"
            "Ask questions about the indexed document collection. "
            "Sources are cited inline and shown below each answer."
        )

        chatbot = gr.Chatbot(
            label="Conversation",
            height=520,
            show_copy_button=True,
            bubble_full_width=False,
            avatar_images=(None, "https://em-content.zobj.net/source/twitter/376/robot_1f916.png"),
        )

        sources_display = gr.HTML(
            value='<div class="status-bar">Sources will appear here after each answer.</div>',
            label="Retrieved Sources",
        )

        with gr.Row():
            msg_input = gr.Textbox(
                placeholder="Ask a question about your documents...",
                show_label=False,
                scale=9,
                container=False,
            )
            send_btn = gr.Button("Send", variant="primary", scale=1)

        with gr.Row():
            clear_btn = gr.Button("🗑 Clear Chat", size="sm")
            status = gr.Markdown(
                f"*{pipeline.vector_store.collection.count()} chunks indexed "
                f"| Hybrid search (semantic + BM25) "
                f"| Memory: last {pipeline.config.memory_length} exchanges*"
            )

        # -- Event handlers ------------------------------------------------

        def respond(user_message: str, chat_history: list):
            if not user_message.strip():
                return "", chat_history, ""

            # Guard: check if any documents are indexed
            if pipeline.vector_store.collection.count() == 0:
                answer = (
                    "⚠️ No documents have been indexed yet. Please add at least 50 documents "
                    "(PDF, DOCX, HTML, TXT, or MD files) to the `./documents/` folder and "
                    "restart the application."
                )
                chat_history = chat_history + [[user_message, answer]]
                return "", chat_history, '<div class="status-bar">No documents indexed.</div>'

            answer, sources = pipeline.query(user_message)
            chat_history = chat_history + [[user_message, answer]]

            # Format sources as HTML cards
            if sources:
                cards = []
                for i, s in enumerate(sources, 1):
                    src = s["metadata"].get("source", "unknown")
                    title = s["metadata"].get("title", "")
                    score = s.get("hybrid_score", s.get("score", 0))
                    preview = s["text"][:250].replace("\n", " ") + "..."
                    card = (
                        f'<div class="source-card">'
                        f"<strong>[Source {i}]</strong> {src}"
                        f"{f' — <em>{title}</em>' if title else ''}"
                        f" (score: {score:.4f})<br>"
                        f"<span style='color:#555'>{preview}</span>"
                        f"</div>"
                    )
                    cards.append(card)
                sources_html = "".join(cards)
            else:
                sources_html = '<div class="status-bar">No relevant sources found.</div>'

            return "", chat_history, sources_html

        def clear_chat():
            pipeline.reset_conversation()
            return [], '<div class="status-bar">Conversation cleared. Sources will appear here.</div>'

        # Wire events
        msg_input.submit(respond, [msg_input, chatbot], [msg_input, chatbot, sources_display])
        send_btn.click(respond, [msg_input, chatbot], [msg_input, chatbot, sources_display])
        clear_btn.click(clear_chat, outputs=[chatbot, sources_display])

    return demo

# ---------------------------------------------------------------------------
# 9. Main Entry Point
# ---------------------------------------------------------------------------

def main():
    """Initialize the pipeline, index documents, and launch the UI."""
    log.info("=" * 60)
    log.info("RAG Application Starting")
    log.info("=" * 60)

    # Ensure NLTK data is available
    try:
        sent_tokenize("Hello world.")
    except LookupError:
        nltk.download("punkt_tab", quiet=True)

    # Validate API key
    api_key = os.getenv("OPENAI_API_KEY", "")
    if not api_key:
        log.warning(
            "OPENAI_API_KEY not set. LLM generation will fail. "
            "Set it with: export OPENAI_API_KEY='sk-...'"
        )

    # Create documents directory if needed
    os.makedirs(CFG.documents_dir, exist_ok=True)

    # Initialize pipeline
    pipeline = RAGPipeline(CFG)

    # Index documents (idempotent — skips already-indexed chunks)
    pipeline.index_documents()

    # Build and launch UI
    demo = build_ui(pipeline)
    log.info(f"Launching Gradio on port {CFG.server_port} (share={CFG.share})")
    demo.launch(
        server_name="0.0.0.0",
        server_port=CFG.server_port,
        share=CFG.share,
        show_error=True,
    )


if __name__ == "__main__":
    main()